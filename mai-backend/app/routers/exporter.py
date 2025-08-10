from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import StreamingResponse, PlainTextResponse
from io import BytesIO
from typing import Literal
from app.schemas import GenerateResponse, RequirementItem

from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm

router = APIRouter()

def _md(doc: GenerateResponse) -> str:
    lines = []
    lines.append(f"# {doc.project_name}")
    lines.append("")
    lines.append(doc.summary)
    lines.append("")
    for cat in doc.categories:
        lines.append(f"## {cat}")
        for r in [x for x in doc.requirements if x.category == cat]:
            lines.append(f"- **{r.priority}** {r.text}")
            if r.acceptance_criteria:
                for ac in r.acceptance_criteria:
                    lines.append(f"  - [ ] {ac}")
            if r.rationale:
                lines.append(f"  - _Rationale_: {r.rationale}")
            if r.standard_refs:
                lines.append(f"  - _Standards_: {', '.join(r.standard_refs)}")
        lines.append("")
    return "\n".join(lines)

def _docx(doc: GenerateResponse) -> BytesIO:
    d = Document()
    d.add_heading(doc.project_name, 0)
    p = d.add_paragraph(doc.summary)
    p.style.font.size = Pt(11)

    for cat in doc.categories:
        d.add_heading(cat, level=2)
        for r in [x for x in doc.requirements if x.category == cat]:
            p = d.add_paragraph()
            run = p.add_run(f"{r.priority} — {r.text}")
            run.bold = True
            if r.acceptance_criteria:
                for ac in r.acceptance_criteria:
                    d.add_paragraph(ac, style="List Bullet")
            if r.rationale:
                d.add_paragraph(f"Rationale: {r.rationale}")
            if r.standard_refs:
                d.add_paragraph(f"Standards: {', '.join(r.standard_refs)}")

    bio = BytesIO()
    d.save(bio)
    bio.seek(0)
    return bio

def _pdf(doc: GenerateResponse) -> BytesIO:
    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4

    y = height - 2*cm
    def draw_line(text, bold=False):
        nonlocal y
        if y < 2*cm:
            c.showPage(); y = height - 2*cm
        if bold:
            c.setFont("Helvetica-Bold", 11)
        else:
            c.setFont("Helvetica", 10)
        c.drawString(2*cm, y, text[:110])
        y -= 14

    c.setFont("Helvetica-Bold", 16)
    c.drawString(2*cm, y, doc.project_name); y -= 22
    c.setFont("Helvetica", 10)
    for line in doc.summary.splitlines():
        draw_line(line)

    for cat in doc.categories:
        y -= 6
        draw_line(cat, bold=True)
        for r in [x for x in doc.requirements if x.category == cat]:
            draw_line(f"{r.priority} — {r.text}")
            for ac in r.acceptance_criteria or []:
                draw_line(f"• {ac}")
            if r.rationale:
                draw_line(f"Rationale: {r.rationale}")
            if r.standard_refs:
                draw_line(f"Standards: {', '.join(r.standard_refs)}")

    c.save()
    bio.seek(0)
    return bio

@router.post("/export")
def export_file(doc: GenerateResponse, format: Literal["pdf","docx","md"] = Query("pdf")):
    try:
        if format == "md":
            text = _md(doc)
            return PlainTextResponse(text, media_type="text/markdown")
        elif format == "docx":
            bio = _docx(doc)
            return StreamingResponse(bio, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     headers={"Content-Disposition": f'attachment; filename="{doc.project_name}.docx"'})
        elif format == "pdf":
            bio = _pdf(doc)
            return StreamingResponse(bio, media_type="application/pdf",
                                     headers={"Content-Disposition": f'attachment; filename="{doc.project_name}.pdf"'})
        else:
            raise HTTPException(status_code=400, detail="Unsupported format")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
